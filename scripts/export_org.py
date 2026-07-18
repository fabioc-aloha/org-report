"""Export the Org Directory canvas state to a polished Word document.

Usage:
    python export_org.py
    python export_org.py --state PATH --docx-out PATH --pdf-out PATH

Defaults read/write the org-directory canvas artifacts so the canvas refresh works
without any flags. Pass --pdf-out '' to skip PDF conversion.
"""
import argparse, json, os, re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

_DEFAULT_STATE = os.path.expanduser(r"~\.copilot\extensions\org-directory\artifacts\directory.json")
_DEFAULT_PDF   = os.path.expanduser(r"~\.copilot\extensions\org-directory\artifacts\report.pdf")

_argp = argparse.ArgumentParser(description="Export org-directory state to DOCX + PDF")
_argp.add_argument("--state", default=_DEFAULT_STATE, help="Path to directory.json state file")
_argp.add_argument("--docx-out", default=None, help="Override DOCX output path (default: <script dir>/<Target>-Org-Report.docx)")
_argp.add_argument("--pdf-out", default=_DEFAULT_PDF, help="PDF output path (empty string skips PDF)")
_args = _argp.parse_args()

STATE = _args.state
_target_slug = None  # set after loading state


# Palette
ACCENT = RGBColor(0x0F, 0x62, 0xFE)   # blue
INK    = RGBColor(0x1F, 0x23, 0x28)   # near-black
MUTED  = RGBColor(0x6A, 0x73, 0x7D)   # gray
BAR    = "0F62FE"                     # hex for XML shading
BG     = "F6F8FA"                     # light gray box
BASE_FONT = "Aptos"

# Team color palette — distinct professional hues; (bar_hex, tint_hex_for_bg)
TEAM_PALETTE = [
    ("0F62FE", "EAF3FF"),  # blue
    ("D97706", "FEF3E2"),  # amber
    ("059669", "E6F6EF"),  # emerald
    ("DC2626", "FDECEC"),  # red
    ("7C3AED", "F1EBFC"),  # violet
    ("0891B2", "E4F4F8"),  # cyan
    ("DB2777", "FCE7F1"),  # pink
    ("65A30D", "F0F7E4"),  # lime
    ("EA580C", "FDEDE0"),  # orange
    ("4F46E5", "ECEBFB"),  # indigo
    ("0D9488", "E1F3F1"),  # teal
    ("B45309", "F9EEDF"),  # bronze
    ("6D28D9", "EEE7FA"),  # deep purple
    ("BE185D", "FADDE8"),  # magenta
    ("15803D", "E4F1E8"),  # forest
    ("1D4ED8", "E1EAFB"),  # royal blue
    ("A16207", "F5EBDA"),  # ochre
    ("9333EA", "F2E6FB"),  # purple
    ("047857", "E0F1E9"),  # dark emerald
    ("BE123C", "F9DBE2"),  # crimson
    ("2563EB", "E3ECFA"),  # bright blue
    ("C2410C", "F9E3D5"),  # rust
]

with open(STATE, "r", encoding="utf-8") as f:
    state = json.load(f)

# Strip AI citation markers like [turn1search34] from any string in the state
_CITE_RE = re.compile(r"\s*\[(?:turn[^\]]*|Sources?:[^\]]*|citation:[^\]]*|ref:[^\]]*)\]", re.IGNORECASE)
def _scrub(v):
    if isinstance(v, str):
        return _CITE_RE.sub("", v)
    if isinstance(v, list):
        return [_scrub(x) for x in v]
    if isinstance(v, dict):
        return {k: _scrub(x) for k, x in v.items()}
    return v
state = _scrub(state)

profiles = state.get("profiles", {})
chart = state.get("orgChart") or {}
chain = chart.get("managementChain") or []
target = chart.get("target")
directs = chart.get("directReports") or []

# Filename derived from target
def _slugify(name):
    return "".join(ch if ch.isalnum() else "" for ch in (name or "Org"))
OUT = _args.docx_out or os.path.join(os.path.dirname(__file__), f"{_slugify(target.get('displayName') if target else 'Org')}-Org-Report.docx")


def key_of(name): return (name or "").strip().lower()

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# --- Base style ---
normal = doc.styles["Normal"]
normal.font.name = BASE_FONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(2)

# Configure heading styles
def _style(name, size, color=INK, bold=True, space_before=12, space_after=6):
    s = doc.styles[name]
    s.font.name = BASE_FONT
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = color
    pf = s.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.keep_with_next = True

_style("Heading 1", 20, ACCENT, True, 24, 8)
_style("Heading 2", 14, INK, True, 18, 6)
_style("Heading 3", 12, INK, True, 12, 4)

# Footer page numbers
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
run.font.size = Pt(9)
run.font.color.rgb = MUTED
fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText'); instr.text = 'PAGE'
fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)

# ---------- Helpers ----------
def add_para(text=None, *, style=None, size=None, color=None, bold=False, italic=False,
             align=None, space_before=None, space_after=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None: p.alignment = align
    if space_before is not None: p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        if bold: r.bold = True
        if italic: r.italic = True
    return p

def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def cell_border(cell, edges=("top","bottom","left","right"), size=4, color="D0D7DE"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders'); tc_pr.append(tc_borders)
    for edge in edges:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:color'), color)
        tc_borders.append(el)

def hr(color="D0D7DE"):
    """Horizontal rule as a bottom-border paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

def accent_bar_heading(text, level=2, bar_hex=None, subtitle=None):
    """Section heading with left color bar via a 1-row table."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(0.35)
    tbl.columns[1].width = Cm(16.5)
    bar, cell = tbl.rows[0].cells
    bar.width = Cm(0.35); cell.width = Cm(16.5)
    shade_cell(bar, bar_hex or BAR)
    bar.text = ""
    # remove all borders
    for c in (bar, cell):
        c._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for e in ("top","bottom","left","right","insideH","insideV"):
            b = OxmlElement(f'w:{e}'); b.set(qn('w:val'), 'nil'); borders.append(b)
        c._tc.tcPr.append(borders)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = BASE_FONT
    r.font.size = Pt(16 if level == 2 else 13)
    r.bold = True
    r.font.color.rgb = INK
    if subtitle:
        sp = cell.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(0)
        sr = sp.add_run(subtitle)
        sr.font.name = BASE_FONT
        sr.font.size = Pt(10)
        sr.italic = True
        sr.font.color.rgb = MUTED
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def tag_line(label, items):
    if not items: return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}:  "); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = MUTED
    r2 = p.add_run("  ".join(f"·  {it}" for it in items))
    r2.font.size = Pt(10)

def bullet_list(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs: r.font.size = Pt(11)
        if not p.runs:
            p.add_run(it).font.size = Pt(11)
        else:
            p.runs[0].text = it

# ============ COVER ============
for _ in range(3): doc.add_paragraph()
add_para("ORGANIZATION REPORT", size=11, color=ACCENT, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
if target:
    add_para(target.get("displayName",""), size=32, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(target.get("jobTitle",""), size=14, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(target.get("department",""), size=11, color=MUTED, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
hr()
add_para(f"Generated {datetime.now().strftime('%B %d, %Y at %H:%M')}",
         size=10, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
_vendor_count = sum(1 for pr in profiles.values() if "(" in (pr.get("displayName") or ""))
_fte_count = len(profiles) - _vendor_count
add_para(f"{_fte_count} FTE profiles  ·  {_vendor_count} vendor profiles",
         size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

# Byline + methodology
add_para("Prepared by Fabio Correa", size=11, bold=True, color=INK,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("GCX Data & Insights  ·  Microsoft", size=9, color=MUTED, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

_method = (
    "This brief profiles the leadership chain, direct reports, extended teams, and vendor "
    "partners in the target executive's organization. The reporting structure comes from "
    "Microsoft Graph, and each profile is synthesized by an AI agent from the author's "
    "accessible Microsoft 365 signals (mail, calendar, Teams, and documents) using GitHub "
    "Copilot CLI with the WorkIQ connector, then typeset through a Python export pipeline. "
    "Treat it as a point-in-time working snapshot, not an authoritative HR record."
)
mp = doc.add_paragraph(_method)
mp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
mp.paragraph_format.left_indent = Cm(2.2); mp.paragraph_format.right_indent = Cm(2.2)
mp.paragraph_format.line_spacing = 1.15; mp.paragraph_format.space_after = Pt(8)

# Disclosure on data access limits
_disclosure = (
    "Disclosure: Profiles draw only on the Microsoft 365 signals the author is authorized "
    "to see. Anything protected by information barriers, sensitivity labels, or Purview "
    "policies is invisible to this report, so collaborations and areas of ownership may be "
    "under-represented. Summaries are AI-generated and can contain errors or misattributed "
    "activity, so verify before acting on any specific claim. This document contains "
    "Microsoft confidential personnel information; handle it accordingly and do not "
    "redistribute."
)
dp = doc.add_paragraph()
dr = dp.add_run(_disclosure)
dr.italic = True; dr.font.size = Pt(9); dr.font.color.rgb = MUTED
dp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp.paragraph_format.left_indent = Cm(2.2); dp.paragraph_format.right_indent = Cm(2.2)
dp.paragraph_format.line_spacing = 1.15; dp.paragraph_format.space_after = Pt(0)
for r in mp.runs:
    r.font.size = Pt(9); r.font.color.rgb = MUTED; r.italic = True

doc.add_page_break()

# ============ MANAGEMENT CHAIN ============
accent_bar_heading("Management chain")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
entries = list(chain) + ([target] if target else [])
for i, entry in enumerate(entries):
    if i:
        arrow = p.add_run("   →   "); arrow.font.color.rgb = ACCENT; arrow.bold = True
    r = p.add_run(entry.get("displayName",""))
    r.bold = True
    if entry is target: r.font.color.rgb = ACCENT

# Chain table
tbl = doc.add_table(rows=len(entries), cols=2)
tbl.autofit = False
tbl.columns[0].width = Cm(5.5); tbl.columns[1].width = Cm(11.5)
for i, entry in enumerate(entries):
    is_target = entry is target
    left, right = tbl.rows[i].cells
    left.width = Cm(5.5); right.width = Cm(11.5)
    p_l = left.paragraphs[0]
    r = p_l.add_run(("⭐  " if is_target else "") + entry.get("displayName",""))
    r.bold = True; r.font.size = Pt(11)
    if is_target: r.font.color.rgb = ACCENT
    p_r = right.paragraphs[0]
    r2 = p_r.add_run(entry.get("jobTitle",""))
    r2.font.size = Pt(10)
    if entry.get("department"):
        r3 = p_r.add_run(f"  ·  {entry['department']}")
        r3.font.size = Pt(10); r3.font.color.rgb = MUTED
    for c in (left, right):
        cell_border(c, edges=("bottom",), color="E4E7EB")
    if is_target:
        shade_cell(left, "EAF3FF")
        shade_cell(right, "EAF3FF")

# ============ DIRECT REPORTS TABLE ============
doc.add_paragraph()
accent_bar_heading(f"Direct reports  ({len(directs)})")
tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.autofit = False
widths = [Cm(4.6), Cm(4.4), Cm(4.8), Cm(3.2)]
for col, w in zip(tbl.columns, widths): col.width = w

hdrs = ["Name", "Title", "Department", "Sub-reports"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    cell.width = widths[i]
    shade_cell(cell, BAR)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(h.upper()); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    cell_border(cell, color="0F62FE")

for idx, r in enumerate(directs):
    row = tbl.add_row()
    for i, val in enumerate([r.get("displayName",""), r.get("jobTitle",""), r.get("department","")]):
        c = row.cells[i]; c.width = widths[i]
        p = c.paragraphs[0]
        run = p.add_run(val); run.font.size = Pt(10)
        if i == 0: run.bold = True
        cell_border(c, color="E4E7EB")
    subs = r.get("reports") or []
    c = row.cells[3]; c.width = widths[3]
    p = c.paragraphs[0]
    if subs:
        run = p.add_run("\n".join(s.get("displayName","") for s in subs))
        run.font.size = Pt(9); run.font.color.rgb = MUTED
    else:
        run = p.add_run("—"); run.font.size = Pt(9); run.font.color.rgb = MUTED
    cell_border(c, color="E4E7EB")
    if idx % 2 == 1:
        for c in row.cells: shade_cell(c, "FAFBFC")

doc.add_page_break()

def render_profile(entry, profile, is_first=False, team_hex=None, team_tint=None):
    tint = team_tint or "F6F8FA"
    hexc = team_hex or BAR
    if not is_first: doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # Name row: color chip + name in a 1-row table (keeps chip flush with name)
    ntbl = doc.add_table(rows=1, cols=2)
    ntbl.autofit = False
    ntbl.columns[0].width = Cm(0.25); ntbl.columns[1].width = Cm(16.5)
    chip, ncell = ntbl.rows[0].cells
    chip.width = Cm(0.25); ncell.width = Cm(16.5)
    shade_cell(chip, hexc)
    chip.text = ""
    for c in (chip, ncell):
        c._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for e in ("top","bottom","left","right","insideH","insideV"):
            b = OxmlElement(f'w:{e}'); b.set(qn('w:val'), 'nil'); borders.append(b)
        c._tc.tcPr.append(borders)
    ncell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    np = ncell.paragraphs[0]
    np.paragraph_format.space_before = Pt(4); np.paragraph_format.space_after = Pt(0)
    nr = np.add_run("  " + entry.get("displayName","(unknown)"))
    nr.bold = True; nr.font.size = Pt(18); nr.font.color.rgb = INK

    # Title / dept line
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(2)
    r = meta_p.add_run(entry.get("jobTitle","")); r.font.size = Pt(11); r.italic = True
    if entry.get("department"):
        r2 = meta_p.add_run(f"   ·   {entry['department']}"); r2.font.size = Pt(11); r2.italic = True; r2.font.color.rgb = MUTED
    # Contact line
    contact_p = doc.add_paragraph()
    contact_p.paragraph_format.space_after = Pt(6)
    bits = []
    if entry.get("email"): bits.append(entry["email"])
    if entry.get("officeLocation"): bits.append(entry["officeLocation"])
    if bits:
        r = contact_p.add_run("   ·   ".join(bits))
        r.font.size = Pt(9); r.font.color.rgb = MUTED
    hr(color=hexc)

    if not profile:
        add_para("No enriched WorkIQ profile available.", size=10, italic=True, color=MUTED, space_after=12)
        return

    if profile.get("summary"):
        p = doc.add_paragraph(profile["summary"])
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs: run.font.size = Pt(11)

    work = profile.get("recentWork") or []
    if work:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(4); h.paragraph_format.space_after = Pt(2)
        r = h.add_run("RECENT WORK"); r.bold = True; r.font.size = Pt(9)
        # Convert team hex to RGBColor for the label
        r.font.color.rgb = RGBColor(int(hexc[0:2],16), int(hexc[2:4],16), int(hexc[4:6],16))
        for it in work:
            b = doc.add_paragraph(it, style="List Bullet")
            b.paragraph_format.space_after = Pt(1)
            for run in b.runs: run.font.size = Pt(10.5)

    for label, items in [("Topics", profile.get("topics")),
                          ("Expertise", profile.get("expertise")),
                          ("Key collaborators", profile.get("collaborators"))]:
        tag_line(label, items or [])

    if profile.get("lastRefreshed"):
        try:
            ts = datetime.fromisoformat(profile["lastRefreshed"].replace("Z","+00:00"))
            when = ts.strftime("%B %d, %Y")
        except Exception:
            when = profile["lastRefreshed"]
        add_para(f"Profile refreshed {when}", size=8, italic=True, color=MUTED, space_before=8)

# ---- Target profile first ----
first_rendered = False

# ============ PROFILES ============
accent_bar_heading("Individual profiles")

def _flatten_team(direct):
    out = []
    for s in (direct.get("reports") or []):
        out.append(s)
        for ss in (s.get("reports") or []):
            out.append(ss)
    return out

def _name_key(entry):
    return (entry.get("displayName") or "").strip().lower()

def _is_vendor(entry):
    email = (entry.get("email") or "").lower()
    name = entry.get("displayName") or ""
    if email.startswith("v-"): return True
    if "(" in name and ")" in name:
        inside = name[name.index("(")+1:name.rindex(")")]
        if inside.strip() and "NON EA" not in inside.upper():
            return True
    return False

# Split Nicole's directs into three buckets
team_leads = []      # FTE directs WITH reports
direct_ics = []      # FTE directs WITHOUT reports (individual contributors)
direct_vendors = []  # vendor directs (go straight to appendix)
for d in directs:
    if _is_vendor(d):
        direct_vendors.append(d)
    elif d.get("reports"):
        team_leads.append(d)
    else:
        direct_ics.append(d)

team_leads.sort(key=_name_key)
direct_ics.sort(key=_name_key)
direct_vendors.sort(key=_name_key)

# Color assignment: only real teams get palette colors
IC_COLOR = ("475569", "F1F5F9")  # slate for ICs
team_color = {}
for i, d in enumerate(team_leads):
    team_color[_name_key(d)] = TEAM_PALETTE[i % len(TEAM_PALETTE)]

# Precompute FTE / vendor split per team
team_ftes = {}
team_vendors = {}
for d in team_leads:
    members = _flatten_team(d)
    key = _name_key(d)
    team_ftes[key] = [m for m in members if not _is_vendor(m)]
    team_vendors[key] = [m for m in members if _is_vendor(m)]
total_team_vendors = sum(len(v) for v in team_vendors.values())
total_vendors_in_appendix = total_team_vendors + len(direct_vendors)

# ---- Team index page ----
add_para("Teams at a glance", size=14, bold=True, color=INK, space_before=8, space_after=8)
itbl = doc.add_table(rows=1, cols=5)
itbl.autofit = False
iwidths = [Cm(0.6), Cm(5.8), Cm(6.0), Cm(2.1), Cm(2.1)]
for col, w in zip(itbl.columns, iwidths): col.width = w
hdr_cells = itbl.rows[0].cells
for i, h in enumerate(["", "TEAM LEAD", "TITLE", "FTEs", "VENDORS"]):
    c = hdr_cells[i]; c.width = iwidths[i]
    shade_cell(c, BAR)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(h); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    cell_border(c, color="0F62FE")
for d in team_leads:
    row = itbl.add_row()
    hexc, tint = team_color[_name_key(d)]
    key = _name_key(d)
    n_fte = 1 + len(team_ftes[key])
    n_vend = len(team_vendors[key])
    c0 = row.cells[0]; c0.width = iwidths[0]
    shade_cell(c0, hexc); c0.text = ""
    cell_border(c0, color=hexc)
    c1 = row.cells[1]; c1.width = iwidths[1]
    p = c1.paragraphs[0]
    r = p.add_run(d.get("displayName","")); r.bold = True; r.font.size = Pt(10.5)
    cell_border(c1, color="E4E7EB")
    c2 = row.cells[2]; c2.width = iwidths[2]
    p = c2.paragraphs[0]
    r = p.add_run(d.get("jobTitle","")); r.font.size = Pt(10); r.font.color.rgb = MUTED
    cell_border(c2, color="E4E7EB")
    for idx, val in [(3, n_fte), (4, n_vend)]:
        c = row.cells[idx]; c.width = iwidths[idx]
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(val) if val else "—"); r.font.size = Pt(11); r.bold = (val > 0)
        if idx == 4 and val: r.font.color.rgb = MUTED
        cell_border(c, color="E4E7EB")
# IC + direct vendor summary row (as a subtle note under the table)
note_bits = []
if direct_ics: note_bits.append(f"{len(direct_ics)} individual contributor" + ("s" if len(direct_ics) != 1 else "") + " reporting directly")
if direct_vendors: note_bits.append(f"{len(direct_vendors)} direct vendor" + ("s" if len(direct_vendors) != 1 else "") + " (in appendix)")
if note_bits:
    add_para("Plus  " + "  ·  ".join(note_bits), size=10, italic=True, color=MUTED, space_before=8)

doc.add_page_break()

# ---- Leadership chain (management above target) ----
LEADER_COLOR = ("B45309", "F9EEDF")  # bronze/gold for executive leaders
if chain:
    leader_hex, leader_tint = LEADER_COLOR
    subtitle = f"{len(chain)} executive" + ("s" if len(chain) != 1 else "") + f"  ·  reporting chain above {target.get('displayName','') if target else 'the target'}"
    accent_bar_heading("Leadership chain", bar_hex=leader_hex, subtitle=subtitle)
    for i, leader in enumerate(chain):
        if i > 0:
            doc.add_page_break()
        profile = profiles.get(key_of(leader.get("displayName")))
        render_profile(leader, profile, is_first=True, team_hex=leader_hex, team_tint=leader_tint)
    doc.add_page_break()

# ---- Target profile ----
if target:
    profile = profiles.get(key_of(target.get("displayName")))
    render_profile(target, profile, is_first=True, team_hex=BAR, team_tint="EAF3FF")

# ---- Each team: FTEs only ----
for direct in team_leads:
    lead_name = direct.get("displayName", "")
    key = _name_key(direct)
    hexc, tint = team_color[key]
    ftes = team_ftes[key]
    n_vend = len(team_vendors[key])
    doc.add_page_break()
    subtitle_parts = [f"{1 + len(ftes)} FTEs"]
    if n_vend:
        subtitle_parts.append(f"{n_vend} vendor" + ("s" if n_vend != 1 else "") + " in appendix")
    subtitle_parts.append(f"led by {lead_name}")
    subtitle = "  ·  ".join(subtitle_parts)
    accent_bar_heading(f"{lead_name}'s team", bar_hex=hexc, subtitle=subtitle)
    lead_profile = profiles.get(key_of(direct.get("displayName")))
    render_profile(direct, lead_profile, is_first=True, team_hex=hexc, team_tint=tint)
    for member in sorted(ftes, key=_name_key):
        profile = profiles.get(key_of(member.get("displayName")))
        render_profile(member, profile, is_first=False, team_hex=hexc, team_tint=tint)

# ---- Individual contributors reporting to Nicole ----
if direct_ics:
    ic_hex, ic_tint = IC_COLOR
    doc.add_page_break()
    subtitle = f"{len(direct_ics)} individual contributor" + ("s" if len(direct_ics) != 1 else "") + f"  ·  reporting directly to {target.get('displayName','') if target else 'the target'}"
    accent_bar_heading("Individual contributors", bar_hex=ic_hex, subtitle=subtitle)
    first = True
    for ic in direct_ics:
        profile = profiles.get(key_of(ic.get("displayName")))
        render_profile(ic, profile, is_first=first, team_hex=ic_hex, team_tint=ic_tint)
        first = False

# ============ APPENDIX: CONTRACTORS & VENDORS ============
if total_vendors_in_appendix:
    doc.add_page_break()
    add_para("APPENDIX", size=10, bold=True, color=ACCENT,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para("Contractors & Vendors", size=24, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, color=INK)
    n_scopes = sum(1 for k in team_vendors if team_vendors[k]) + (1 if direct_vendors else 0)
    add_para(f"{total_vendors_in_appendix} contractor and vendor profiles across {n_scopes} teams",
             size=11, italic=True, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    hr()

    # Direct-to-target vendors first
    if direct_vendors:
        ic_hex, ic_tint = IC_COLOR
        doc.add_page_break()
        subtitle = f"{len(direct_vendors)} contractor" + ("s" if len(direct_vendors) != 1 else "") + f"  ·  reporting directly to {target.get('displayName','') if target else 'the target'}"
        accent_bar_heading("Reporting directly", bar_hex=ic_hex, subtitle=subtitle)
        first = True
        for member in direct_vendors:
            profile = profiles.get(key_of(member.get("displayName")))
            render_profile(member, profile, is_first=first, team_hex=ic_hex, team_tint=ic_tint)
            first = False

    for direct in team_leads:
        key = _name_key(direct)
        vendors = team_vendors[key]
        if not vendors: continue
        hexc, tint = team_color[key]
        doc.add_page_break()
        subtitle = f"{len(vendors)} contractor" + ("s" if len(vendors) != 1 else "") + f"  ·  under {direct.get('displayName','')}"
        accent_bar_heading(f"{direct.get('displayName','')}'s contractors", bar_hex=hexc, subtitle=subtitle)
        first = True
        for member in sorted(vendors, key=_name_key):
            profile = profiles.get(key_of(member.get("displayName")))
            render_profile(member, profile, is_first=first, team_hex=hexc, team_tint=tint)
            first = False

doc.save(OUT)
print(OUT)

# Also convert to PDF for canvas viewer (skip when --pdf-out '')
if _args.pdf_out:
    try:
        from docx2pdf import convert
        convert(OUT, _args.pdf_out)
        print(f"PDF: {_args.pdf_out}")
    except Exception as e:
        print(f"PDF conversion skipped: {e}")
